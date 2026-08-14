"""
Extracts structured candidate fields from an uploaded resume so apply forms
can be auto-filled (candidate/agency/HR then corrects before submitting).

Text extraction: pypdf for PDF, python-docx for DOCX.
Field extraction: Azure OpenAI (same httpx pattern as ai_rejection_service);
falls back to regex-only extraction (email/phone/links) when AI is unavailable.
"""
import io
import re
import logging
from typing import Any

logger = logging.getLogger(__name__)

MAX_TEXT_CHARS = 15000  # plenty for a resume, keeps the prompt bounded
URL_FIELDS = ("linkedin_url", "github_url", "portfolio_url")


def _with_scheme(url: str) -> str:
    """A resume's visible text often just prints "linkedin.com/in/x" with no
    scheme — stored as-is, that renders as a working-looking link that the
    browser resolves relative to the careers portal's own origin instead of
    navigating to LinkedIn. The real-hyperlink path (_classify_links) doesn't
    need this since a PDF/DOCX anchor target already carries a full URL."""
    return url if url.startswith(("http://", "https://")) else f"https://{url}"

PARSED_FIELDS = [
    "full_name", "email", "phone", "current_location", "total_experience",
    "current_company", "current_designation", "education", "skills",
    "linkedin_url", "github_url", "portfolio_url",
]


def extract_text(content: bytes, content_type: str, filename: str = "") -> str:
    """Best-effort plain-text extraction. Returns '' when the format is unsupported."""
    name = (filename or "").lower()
    try:
        if content_type == "application/pdf" or name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages)

        if (
            content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or name.endswith(".docx")
        ):
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(parts)
    except Exception as exc:
        logger.warning(f"Resume text extraction failed ({content_type}): {exc}")

    return ""


def extract_hyperlinks(content: bytes, content_type: str, filename: str = "") -> list[str]:
    """The actual hyperlink targets embedded in the document (PDF link annotations /
    DOCX hyperlink relationships) — NOT the visible/anchor text.

    Resumes commonly hyperlink a friendly label ("LinkedIn", "Portfolio", or a
    shortened/stale display string) to the real URL. Plain text extraction only
    sees the label, so a regex over extract_text() can silently capture that
    display text (or miss the link entirely) instead of the real href. This
    walks the document's actual link objects so the real URL is used.
    """
    name = (filename or "").lower()
    links: list[str] = []
    try:
        if content_type == "application/pdf" or name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                annots = page.get("/Annots")
                if not annots:
                    continue
                for annot_ref in annots:
                    try:
                        annot = annot_ref.get_object()
                        if annot.get("/Subtype") != "/Link":
                            continue
                        action = annot.get("/A")
                        if action and action.get("/URI"):
                            links.append(str(action["/URI"]))
                    except Exception:
                        continue

        elif (
            content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or name.endswith(".docx")
        ):
            from docx import Document
            doc = Document(io.BytesIO(content))
            for rel in doc.part.rels.values():
                if rel.reltype.endswith("/hyperlink") and rel.is_external:
                    links.append(rel.target_ref)
    except Exception as exc:
        logger.warning(f"Hyperlink extraction failed ({content_type}): {exc}")

    # De-dupe, preserve order
    seen = set()
    unique = []
    for link in links:
        if link not in seen:
            seen.add(link)
            unique.append(link)
    return unique


def _classify_links(links: list[str]) -> dict[str, str | None]:
    """Match extracted hyperlink targets to the linkedin/github/portfolio slots."""
    result: dict[str, str | None] = {"linkedin_url": None, "github_url": None, "portfolio_url": None}
    for link in links:
        low = link.lower()
        if "linkedin.com" in low and not result["linkedin_url"]:
            result["linkedin_url"] = link
        elif "github.com" in low and not result["github_url"]:
            result["github_url"] = link
        elif not any(d in low for d in ("mailto:", "linkedin.com", "github.com")) and not result["portfolio_url"]:
            result["portfolio_url"] = link
    return result


def _regex_fallback(text: str, links: list[str] | None = None) -> dict[str, Any]:
    result: dict[str, Any] = {k: None for k in PARSED_FIELDS}

    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if email:
        result["email"] = email.group(0)

    phone = re.search(r"(?:\+91[\s-]?)?[6-9]\d{9}|\+\d{1,3}[\s-]?\d{6,12}", text)
    if phone:
        result["phone"] = phone.group(0).strip()

    linkedin = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-./]+", text, re.I)
    if linkedin:
        result["linkedin_url"] = _with_scheme(linkedin.group(0))

    github = re.search(r"(?:https?://)?(?:www\.)?github\.com/[\w\-./]+", text, re.I)
    if github:
        result["github_url"] = _with_scheme(github.group(0))

    # Real hyperlink targets (if any were extracted) are authoritative over text
    # regex matches — a resume can hyperlink a friendly label to the real URL,
    # which the text regex above would never see correctly.
    if links:
        for key, val in _classify_links(links).items():
            if val:
                result[key] = val

    result["is_ai_parsed"] = False
    return result


def _build_prompt(text: str, links: list[str] | None = None) -> str:
    links_block = ""
    if links:
        links_block = (
            "\n\nHyperlinks found embedded in the document (these are the real URLs — "
            "prefer these over anything that merely looks like a URL in the visible text):\n"
            + "\n".join(f"- {link}" for link in links)
        )
    return f"""You are a resume parser for a careers portal. Extract candidate details from the resume text below.

Resume text:
\"\"\"
{text[:MAX_TEXT_CHARS]}
\"\"\"{links_block}

Return a JSON object with exactly these keys (use null when the resume does not state a value — never guess or invent):
{{
  "full_name": "candidate's full name",
  "email": "email address",
  "phone": "phone number",
  "current_location": "city/country the candidate is based in",
  "total_experience": "total professional experience, e.g. '5 years'",
  "current_company": "most recent employer ('Fresher' if none)",
  "current_designation": "most recent job title",
  "education": "highest/most recent qualification, e.g. 'B.Tech, CSE, IIT Delhi'",
  "skills": "comma-separated list of key skills",
  "linkedin_url": "LinkedIn profile URL",
  "github_url": "GitHub profile URL",
  "portfolio_url": "personal website/portfolio URL"
}}

Rules:
- Every value must be a string or null. No nested objects or arrays.
- Copy values from the resume; do not fabricate anything that is not present.
- Return only valid JSON, no markdown fences."""


async def parse_resume(content: bytes, content_type: str, filename: str = "") -> dict[str, Any]:
    """
    Returns dict with PARSED_FIELDS keys plus 'is_ai_parsed'.
    Never raises — degrades to regex extraction, then to empty fields.
    """
    from app.config import settings

    text = extract_text(content, content_type, filename)
    links = extract_hyperlinks(content, content_type, filename)
    if not text.strip():
        empty = {k: None for k in PARSED_FIELDS}
        empty.update({k: v for k, v in _classify_links(links).items() if v})
        return {**empty, "is_ai_parsed": False}

    if settings.AZURE_OPENAI_ENDPOINT and settings.AZURE_OPENAI_API_KEY and settings.AZURE_OPENAI_DEPLOYMENT:
        try:
            import httpx, json

            url = (
                f"{settings.AZURE_OPENAI_ENDPOINT.rstrip('/')}/openai/deployments/"
                f"{settings.AZURE_OPENAI_DEPLOYMENT}/chat/completions"
                f"?api-version={settings.AZURE_OPENAI_API_VERSION}"
            )
            payload = {
                "messages": [{"role": "user", "content": _build_prompt(text, links)}],
                "temperature": 0,
                "max_tokens": 800,
                "response_format": {"type": "json_object"},
            }
            async with httpx.AsyncClient(timeout=30) as client:
                resp = await client.post(
                    url,
                    json=payload,
                    headers={"api-key": settings.AZURE_OPENAI_API_KEY},
                )
                resp.raise_for_status()
                raw = resp.json()["choices"][0]["message"]["content"]
                parsed = json.loads(raw)

            result = {}
            for key in PARSED_FIELDS:
                val = parsed.get(key)
                val = val.strip() if isinstance(val, str) and val.strip() else None
                if val and key in URL_FIELDS:
                    val = _with_scheme(val)
                result[key] = val

            # Override with the real hyperlink targets when we found any — more
            # reliable than anything the model read off of visible anchor text.
            for key, val in _classify_links(links).items():
                if val:
                    result[key] = val

            result["is_ai_parsed"] = True
            return result
        except Exception as exc:
            logger.warning(f"Azure OpenAI resume parsing failed, falling back to regex: {exc}")

    return _regex_fallback(text, links)
